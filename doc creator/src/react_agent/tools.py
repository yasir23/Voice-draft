"""This module provides example tools for web scraping and search functionality.

It includes a basic Tavily search function (as an example)

These tools are intended as free examples to get started. For production use,
consider implementing more robust and specialized tools tailored to your needs.
"""

from typing import Any, Callable, List, Optional, cast

from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_core.runnables import RunnableConfig
from langchain_core.tools import InjectedToolArg
from typing_extensions import Annotated

from react_agent.configuration import Configuration



from typing import Any, Callable, List, Optional
from pydantic import BaseModel
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# Define a function to create and format a Word document
async def create_word_doc(
    content: str, file_name: str = "draft.docx"
) -> Optional[str]:
    """Create a Word document with the given content and formatting."""
    try:
        # Create a new Word document
        doc = Document()

        # Add a title
        title = doc.add_heading("Draft Document", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add content
        paragraph = doc.add_paragraph(content)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.style.font.size = Pt(12)

        # Save the document
        file_path = f"./{file_name}"
        doc.save(file_path)
        return file_path
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return None

# Add the new tool to the TOOLS list
# TOOLS: List[Callable[..., Any]] = [create_word_doc]





async def search(
    query: str, *, config: Annotated[RunnableConfig, InjectedToolArg]
) -> Optional[list[dict[str, Any]]]:
    """Search for general web results.
    This function performs a search using the Tavily search engine, which is designed
    to provide comprehensive, accurate, and trusted results. It's particularly useful
    for answering questions about current events.
    """
    configuration = Configuration.from_runnable_config(config)
    wrapped = TavilySearchResults(max_results=configuration.max_search_results)
    result = await wrapped.ainvoke({"query": query})
    return cast(list[dict[str, Any]], result)

# Define the TOOLS list with both tools
TOOLS: List[Callable[..., Any]] = [create_word_doc, search]